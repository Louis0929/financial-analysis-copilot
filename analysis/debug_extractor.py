"""
Debug tool to check what we're actually extracting from DOCX files
"""

def debug_docx_extraction(filepath):
    """Debug DOCX extraction to see what we're getting"""
    try:
        import docx
        doc = docx.Document(filepath)
        
        print("=== DEBUG DOCX EXTRACTION ===")
        print(f"File: {filepath}")
        print(f"Number of paragraphs: {len(doc.paragraphs)}")
        print(f"Number of tables: {len(doc.tables)}")
        
        # Check paragraphs for financial keywords
        print("\n=== PARAGRAPHS WITH FINANCIAL KEYWORDS ===")
        financial_keywords = ['revenue', 'income', 'cost', 'profit', 'margin', '245,122', '88,136']
        for i, para in enumerate(doc.paragraphs):
            text = para.text.lower()
            for keyword in financial_keywords:
                if keyword in text:
                    print(f"Paragraph {i}: {para.text[:100]}...")
                    break
        
        # Detailed table analysis
        print("\n=== TABLE ANALYSIS ===")
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx + 1}:")
            print(f"  Rows: {len(table.rows)}")
            print(f"  Columns: {len(table.rows[0].cells) if table.rows else 0}")
            
            # Show first few rows
            for row_idx, row in enumerate(table.rows[:5]):  # First 5 rows only
                cells = [cell.text.strip() for cell in row.cells]
                print(f"  Row {row_idx}: {' | '.join(cells)}")
                
                # Check for financial numbers
                for cell_text in cells:
                    if any(num in cell_text for num in ['245,122', '88,136', '74,114', '171,008']):
                        print(f"    *** FOUND FINANCIAL DATA: {cell_text} ***")
        
        # Extract and analyze content like our main function
        print("\n=== CONTENT EXTRACTION TEST ===")
        content = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            content += f"\n=== TABLE {table_idx + 1} DATA ===\n"
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_text.append(cell_text if cell_text else "")
                content += " | ".join(row_text) + "\n"
            content += "\n"
        
        print(f"Total content length: {len(content)}")
        
        # Search for specific financial figures
        search_terms = ['245,122', '88,136', '74,114', '171,008', 'total revenue', 'net income']
        print("\n=== SEARCHING FOR SPECIFIC TERMS ===")
        for term in search_terms:
            if term.lower() in content.lower():
                print(f"✅ FOUND: {term}")
                # Show context
                start = max(0, content.lower().find(term.lower()) - 50)
                end = min(len(content), content.lower().find(term.lower()) + 100)
                print(f"   Context: ...{content[start:end]}...")
            else:
                print(f"❌ NOT FOUND: {term}")
        
        return content
        
    except Exception as e:
        print(f"Error: {e}")
        return None

if __name__ == "__main__":
    # Test with a sample file
    print("Debug tool for DOCX extraction")