"""
File reader module for Financial Statement Analysis Co-Pilot
Handles reading various file formats: TXT, PDF, DOCX, XLSX, CSV
"""

import os
import pandas as pd
from pathlib import Path

# Import libraries for different file formats
try:
    import PyPDF2
    import docx
    from openpyxl import load_workbook
except ImportError as e:
    print(f"⚠️ Some file format libraries are missing: {e}")

def read_txt_file(filepath):
    """Read content from a TXT file"""
    with open(filepath, 'r', encoding='utf-8') as file:
        return file.read()

def read_pdf_file(filepath):
    """Read content from a PDF file using advanced extraction"""
    try:
        # Try pdfplumber first (better for tables)
        try:
            import pdfplumber
            content = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    # Extract text
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
                    
                    # Extract tables separately
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            # Convert table to readable text
                            for row in table:
                                if row:
                                    content += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                            content += "\n"
            return content
        except ImportError:
            # Fallback to PyPDF2
            import PyPDF2
            content = ""
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
            return content
    except ImportError:
        raise ImportError("pdfplumber or PyPDF2 is required for PDF files. Install with: pip install pdfplumber PyPDF2")

def read_docx_file(filepath):
    """Read content from a Word document including tables with enhanced extraction"""
    try:
        import docx
        doc = docx.Document(filepath)
        content = ""
        
        # DEBUG: Add extraction metadata
        content += f"=== DOCX EXTRACTION DEBUG ===\n"
        content += f"Total paragraphs: {len(doc.paragraphs)}\n"
        content += f"Total tables: {len(doc.tables)}\n"
        content += f"=== END DEBUG ===\n\n"
        
        # Extract paragraphs with more aggressive number detection
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:  # Only add non-empty paragraphs
                content += para_text + "\n"
                # Look for financial numbers in paragraphs too
                if any(char.isdigit() for char in para_text) and (',' in para_text or '$' in para_text):
                    content += f"[FINANCIAL_DATA_PARAGRAPH]: {para_text}\n"
                paragraph_count += 1
        
        content += f"\n=== EXTRACTED {paragraph_count} NON-EMPTY PARAGRAPHS ===\n\n"
        
        # Extract tables with enhanced debugging
        if len(doc.tables) == 0:
            content += "⚠️ WARNING: NO TABLES FOUND IN DOCUMENT\n\n"
        
        for table_idx, table in enumerate(doc.tables):
            content += f"\n=== TABLE {table_idx + 1} DATA ===\n"
            content += f"Table has {len(table.rows)} rows and {len(table.columns)} columns\n"
            
            # Extract all cells, including nested content
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell_idx, cell in enumerate(row.cells):
                    # Get all text from cell, including nested paragraphs
                    cell_content = ""
                    for para in cell.paragraphs:
                        cell_content += para.text.strip() + " "
                    cell_content = cell_content.strip()
                    
                    # Also try direct cell.text as backup
                    if not cell_content:
                        cell_content = cell.text.strip()
                    
                    row_text.append(cell_content if cell_content else "")
                
                # Always add the row, even if some cells are empty
                row_data = " | ".join(row_text)
                content += f"ROW_{row_idx}: {row_data}\n"
                
                # Flag rows with financial data
                if any(char.isdigit() for char in row_data) and (',' in row_data or '$' in row_data):
                    content += f"[FINANCIAL_ROW_{row_idx}]: {row_data}\n"
            
            content += f"=== END TABLE {table_idx + 1} ===\n\n"
            
            # Also extract table in a more readable format for financial data
            if len(table.rows) > 1:  # Has header row
                content += f"--- TABLE {table_idx + 1} FORMATTED ---\n"
                for row_idx, row in enumerate(table.rows):
                    cells = []
                    for cell in row.cells:
                        cell_content = ""
                        for para in cell.paragraphs:
                            cell_content += para.text.strip() + " "
                        cells.append(cell_content.strip())
                    
                    if row_idx == 0:  # Header row
                        content += "HEADERS: " + " | ".join(cells) + "\n"
                    else:  # Data rows
                        content += f"ROW {row_idx}: " + " | ".join(cells) + "\n"
                content += f"--- END FORMATTED TABLE {table_idx + 1} ---\n\n"
        
        # Final debug summary
        content += f"\n=== FINAL EXTRACTION SUMMARY ===\n"
        content += f"Content length: {len(content)} characters\n"
        content += f"Contains 'revenue': {'revenue' in content.lower()}\n"
        content += f"Contains 'income': {'income' in content.lower()}\n"
        content += f"Contains 'cost': {'cost' in content.lower()}\n"
        content += f"Contains commas: {',' in content}\n"
        content += f"Contains dollar signs: {'$' in content}\n"
        content += "=== END SUMMARY ===\n\n"
        
        return content
    except ImportError:
        raise ImportError("python-docx is required for Word files. Install with: pip install python-docx")

def read_excel_file(filepath):
    """Read content from an Excel file"""
    try:
        # Read all sheets and combine content
        excel_file = pd.ExcelFile(filepath)
        content = ""
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(filepath, sheet_name=sheet_name)
            content += f"\n--- SHEET: {sheet_name} ---\n"
            content += df.to_string(index=False) + "\n"
        
        return content
    except ImportError:
        raise ImportError("pandas and openpyxl are required for Excel files. Install with: pip install pandas openpyxl")

def read_csv_file(filepath):
    """Read content from a CSV file"""
    try:
        df = pd.read_csv(filepath)
        return df.to_string(index=False)
    except ImportError:
        raise ImportError("pandas is required for CSV files. Install with: pip install pandas")

def read_report(filepath):
    """
    Read financial report content from various file formats
    Supports: TXT, PDF, DOCX, XLSX, CSV
    
    Args:
        filepath (str): Path to the financial report file
        
    Returns:
        str: Content of the financial report or None if error
    """
    try:
        # Check if file exists
        if not os.path.exists(filepath):
            print(f"❌ Error: File not found at {filepath}")
            return None
        
        # Get file extension
        file_extension = Path(filepath).suffix.lower()
        
        # Read based on file type
        if file_extension == '.txt':
            content = read_txt_file(filepath)
        elif file_extension == '.pdf':
            content = read_pdf_file(filepath)
        elif file_extension == '.docx':
            content = read_docx_file(filepath)
        elif file_extension in ['.xlsx', '.xls']:
            content = read_excel_file(filepath)
        elif file_extension == '.csv':
            content = read_csv_file(filepath)
        else:
            print(f"❌ Unsupported file format: {file_extension}")
            return None
        
        print(f"✅ Successfully loaded {file_extension.upper()} report from: {filepath}")
        print(f"📄 Content length: {len(content)} characters")
        return content
        
    except Exception as e:
        print(f"❌ Error reading file: {e}")
        return None 